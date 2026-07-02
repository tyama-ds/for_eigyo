"""DocQA — 1つの文書を「散文=RAG / 表=TableQA」に自動振り分けして問い合わせる。

同じファイルの中で:
- **本文（散文）** は PagedRAG（ベクトル検索・ページ出典）で扱う
- **表** は抽出して DataFrame 化し TableQA（text-to-pandas）で集計・計算する

`ask()` は質問を分類し、計算/集計/条件抽出 → TableQA、説明/検索 → RAG に自動ルーティング。
`route="rag"` / `route="table"` で強制も可能。

対応: PDF（pdfplumber で表抽出）/ Excel（各シート）/ Word（python-docx）/ CSV。
任意依存: `pip install -e ".[tables,office]"`（pdfplumber / python-docx / openpyxl / pandas）
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from .client import complete
from .pagedrag import PagedRAG
from .tableqa import TableQA

_TABLE_HINTS = ("合計", "平均", "最大", "最小", "件数", "何件", "割合", "比率", "率",
                "上位", "ランキング", "sum", "average", "mean", "count", "total",
                "max", "min", "より大き", "より小さ", "以上", "以下", "集計", "計算")


@dataclass
class DocResult:
    route: str          # "table" | "rag"
    result: object      # TableAnswer or PagedRAG Answer

    @property
    def text(self) -> str:
        return (getattr(self.result, "answer", "") or getattr(self.result, "text", "")
                or str(self.result))

    def __str__(self) -> str:
        return f"[route = {self.route}]\n{self.result}"


class DocQA:
    """1文書を散文=RAG / 表=TableQA に振り分けて問い合わせる。"""

    def __init__(self, path: str | Path, *, storage_dir: str | Path | None = None,
                 title: str | None = None, prose: bool = True):
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"ファイルがありません: {path}")
        self.path = path
        self.title = title or path.stem
        base = Path(storage_dir) if storage_dir else Path("./storage/docqa") / _safe(self.title)

        # 表 → DataFrame 群 → TableQA
        self.tables = _extract_dataframes(path)
        self._tableqa = TableQA(self.tables) if self.tables else None

        # 本文（散文）→ PagedRAG
        self._rag: PagedRAG | None = None
        if prose:
            self._rag = PagedRAG(storage_dir=str(base / "prose"))
            try:
                self._rag.add_book(path, title=self.title)
            except Exception as e:  # noqa: BLE001
                print(f"[DocQA] 本文の索引化に失敗（表のみで続行）: {e}")
                self._rag = None

        print(f"[DocQA] {self.title}: 表 {len(self.tables)} 個 / 本文RAG {'あり' if self._rag else 'なし'}")

    # ---- 問い合わせ ----
    def ask(self, question: str, *, route: str = "auto") -> DocResult:
        if route not in ("auto", "rag", "table"):
            raise ValueError('route は "auto" / "rag" / "table"')
        chosen = self._route(question) if route == "auto" else route

        if chosen == "table":
            if self._tableqa is None:
                if self._rag is None:
                    raise RuntimeError("表も本文も利用できません。")
                chosen = "rag"  # 表が無ければ RAG へフォールバック
            else:
                try:
                    return DocResult(route="table", result=self._tableqa.ask(question))
                except Exception as e:  # noqa: BLE001
                    # 生成コードの拒否/実行失敗時、auto なら本文RAGで答えを試みる
                    if route == "table" or self._rag is None:
                        raise
                    print(f"[DocQA] 表での回答に失敗したため本文RAGへフォールバック: {e}")
                    chosen = "rag"

        if self._rag is None:
            if self._tableqa is not None:
                return DocResult(route="table", result=self._tableqa.ask(question))
            raise RuntimeError("本文RAG が利用できません。")
        return DocResult(route="rag", result=self._rag.query(question, title=self.title))

    def table_names(self) -> list[str]:
        return list(self.tables.keys())

    # ---- ルーティング ----
    def _route(self, question: str) -> str:
        if not self.tables:
            return "rag"
        if self._rag is None:
            return "table"
        # まずヒューリスティック（計算系の語があれば table 寄り）
        hint = any(h in question.lower() for h in _TABLE_HINTS)
        try:
            from .bookindex import llm_json

            res = llm_json(
                "質問が『表データの集計・計算・条件抽出』なら table、『文章の説明・検索・"
                "要約』なら text を選んでください。JSON のみ: {\"route\": \"table|text\"}\n\n"
                f"表の有無: あり（{len(self.tables)}個）\n質問: {question}"
            )
            r = (res or {}).get("route", "") if isinstance(res, dict) else ""
            if r == "table":
                return "table"
            if r == "text":
                return "rag"
        except Exception:  # noqa: BLE001
            pass
        return "table" if hint else "rag"


# --------------------------------------------------------------------------
# 表抽出 → DataFrame 群
# --------------------------------------------------------------------------

def _extract_dataframes(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix in (".xlsx", ".xls"):
        return _dfs_from_excel(path)
    if suffix == ".csv":
        return _dfs_from_csv(path)
    if suffix == ".pdf":
        return _dfs_from_pdf(path)
    if suffix in (".docx", ".doc"):
        return _dfs_from_docx(path)
    return {}


def _rows_to_df(rows: list, name: str):
    import pandas as pd

    rows = [r for r in rows if r is not None]
    if not rows:
        return None
    header = ["" if c is None else str(c) for c in rows[0]]
    # ヘッダらしさ: 全て非空かつ重複が少ない
    if len(set(header)) == len(header) and all(h.strip() for h in header) and len(rows) > 1:
        try:
            return pd.DataFrame(rows[1:], columns=header)
        except Exception:  # noqa: BLE001
            pass
    return pd.DataFrame(rows)


def _dfs_from_excel(path: Path) -> dict:
    try:
        import pandas as pd

        # read_excel は openpyxl を遅延 import するため、ここも try 内に置く
        data = pd.read_excel(path, sheet_name=None)
    except ImportError:
        print("[DocQA] Excel には pandas/openpyxl が必要: pip install openpyxl pandas（表なしで続行）")
        return {}
    except Exception as e:  # noqa: BLE001
        print(f"[DocQA] Excel の読み込みに失敗（表なしで続行）: {e}")
        return {}
    return {str(k): v for k, v in data.items()}


def _dfs_from_csv(path: Path) -> dict:
    import pandas as pd

    return {path.stem: pd.read_csv(path)}


def _dfs_from_pdf(path: Path) -> dict:
    try:
        import pdfplumber
    except ImportError:
        print("[DocQA] PDF の表抽出には pdfplumber が必要: pip install pdfplumber（表なしで続行）")
        return {}
    out = {}
    try:
        with pdfplumber.open(str(path)) as pdf:
            for pno, page in enumerate(pdf.pages, start=1):
                for ti, tb in enumerate(page.extract_tables() or [], start=1):
                    df = _rows_to_df(tb, f"p{pno}_{ti}")
                    if df is not None and not df.empty:
                        out[f"table_p{pno}_{ti}"] = df
    except Exception as e:  # noqa: BLE001
        print(f"[DocQA] PDF 表抽出に失敗（表なしで続行）: {e}")
    return out


def _dfs_from_docx(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        print("[DocQA] Word の表抽出には python-docx が必要: pip install python-docx（表なしで続行）")
        return {}
    out = {}
    try:
        doc = Document(str(path))
        for ti, tb in enumerate(doc.tables, start=1):
            rows = [[cell.text for cell in row.cells] for row in tb.rows]
            df = _rows_to_df(rows, f"t{ti}")
            if df is not None and not df.empty:
                out[f"table_{ti}"] = df
    except Exception as e:  # noqa: BLE001
        print(f"[DocQA] Word 表抽出に失敗（表なしで続行）: {e}")
    return out


def _safe(title: str) -> str:
    """正規化衝突（全角/半角括弧の差など）で別文書の storage を共有しないようハッシュを付与。"""
    import hashlib

    base = re.sub(r"[^\w\-.]+", "_", title)[:100] or "doc"
    return f"{base}_{hashlib.md5(title.encode('utf-8')).hexdigest()[:8]}"
